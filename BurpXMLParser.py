from html.parser import HTMLParser
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup
import csv
import base64
import os
import optparse
import sys
import logging
import logging.config
from logging.config import fileConfig
from docx import Document
from docx.shared import RGBColor
import matplotlib.pyplot as plt
import io
import re



# Should be done better with loghandler. But cheap way to clear the issue log file on run
# Have to delete the log file before the logged is init
def deleteFile(file):
    if os.path.exists(file):
        os.remove(file)


issues_logFile = os.path.join('issues', 'created-issues.log')
deleteFile(issues_logFile)

LOGGING_LEVELS = {'critical': logging.CRITICAL,
                  'error': logging.ERROR,
                  'warning': logging.WARNING,
                  'info': logging.INFO,
                  'debug': logging.DEBUG}

logging.config.fileConfig('logging.conf')

# create logger
logger = logging.getLogger()
status_logger = logging.getLogger('xmlparser.status')
issue_logger = logging.getLogger('xmlparser.issues')

# define globals
global issueList
global vulnList
global paths
global skippedVulnList
global xmlFileIn
global docOutFile
global cli_XMLFILE
global cli_WORDFILE
global cli_CSVFILE
global cli_XMLPROCESSDIR

# define Cli variables for our ARGs
cli_XMLFILE = ""
cli_WORDFILE = ""
cli_CSVFILE = ""
cli_XMLPROCESSDIR = ""
# Set input and output files
xmlFileIn = cli_XMLFILE
docOutFile = cli_WORDFILE

issueList = []
vulnList = []
skippedVulnList = []
paths = []
# init Document
document = Document()

"""
Basic Layout of Word DOc

{TITLE HEADER} ({Risk Level})

{Header 3.text 'Overview'}
{Paragraph with finding overview}

{Header 3.text 'Evidence'}
{Paragraph with finding Evidence}

{SCREENSHOT OR SNIPPET}

{Header 3.text 'Recommendation'}
{Paragraph with Recommendation}


"""


class MLStripper(HTMLParser):

    def __init__(self):
        super().__init__()
        self.reset()
        self.fed = []

    def handle_data(self, d):
        self.fed.append(d)

    def get_data(self):
        return ''.join(self.fed)


def strip_tags(html):
    s = MLStripper()
    s.feed(html)
    return s.get_data()


def buildWordDoc(name, severity, host, ip, path, location, issueBackground, issueDetail, remediationBackground,
                 vulnerabilityClassification):
    # refer to https://python-docx.readthedocs.io/en/latest/
    # we init the doc at the start of this script
    # then save it in the main function after everything is built.
    location = str(location)
    orig_location = location
    loc_count = location.count('/')
    status_logger.debug('Location String {} location count : {}'.format(location, loc_count))
    # this is a word formatting fix. If the location is / then we add host URL.
    if loc_count < 2:
        status_logger.debug('Location/Path is Default "/" ')
        # full_location = os.path.join(host, location)
        full_location = host + location
        location = full_location
    status_logger.debug('Location is Now {}'.format(location))
    # reformat data if needed
    # cheap oncoding of the comma by replacing it with |.
    # probably need to move all these to after we pass data to the word function.
    # then the commas could be fixed for building the CSV.

    issueBackground = str(issueBackground).replace('|', ',')
    remediationBackground = str(remediationBackground)
    # strip HTML tags using our function instead of string replacements inline.
    remediationBackground = strip_tags(remediationBackground)
    severity = str(severity)
    severity = severity + ' Risk'
    # use title to fix Capitals
    severity = severity.title()
    # Build Our header format here.
    if severity == "High Risk":
        color = RGBColor(255, 0, 0)  # Red
    elif severity == "Medium Risk":
        color = RGBColor(255, 165, 0)  # Orange
    elif severity == "Low Risk":
        color = None
    else:
        color = RGBColor(0, 128, 0)  # Green
    
    # Build header with colored text
    build_header = '{} ({})'.format(name, severity)

    if document.paragraphs:  # Check if document contains any paragraphs
        document.add_page_break()
    header = document.add_heading(level=2)
    run = header.add_run(build_header)
    if color:
        font = run.font
        font.color.rgb = color
        font.bold = True  # Make the text bold
    document.add_heading("Vulnerable Host:", level=3)
    paragraph = document.add_paragraph(host)
    document.add_heading("Vulnerable URL:", level=3)
    # fixing location string so HTTP isnt included twice.
    if 'http' in location:
        location = orig_location
    host_url = host + location
    paragraph = document.add_paragraph(host_url)
    document.add_heading("Technical Details:", level=3)
    table = document.add_table(rows=1, cols=2)
    # adjusted cell alignment here manually.
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'IP:'
    hdr_cells[0].width = Inches(.00)
    hdr_cells[1].text = ip
    hdr_cells[1].width = Inches(.5)
    hdr_cells[1].left_margin = .1
    row_cells = table.add_row().cells
    row_cells[0].text = 'Path:'
    row_cells[0].width = Inches(.00)
    row_cells[1].text = path
    row_cells[1].width = Inches(.5)
    row_cells[1].left_margin = .1

    document.add_heading("Overview:", level=3)
    issueBackground = strip_tags(issueBackground)
    paragraph = document.add_paragraph(issueBackground)

    document.add_heading("Evidence:", level=3)
    issueDetail = strip_tags(issueDetail)
    issueDetail = issueDetail.replace('","', "")
    paragraph = document.add_paragraph(issueDetail)

    document.add_heading("Recommendation:", level=3)
    remediationBackground = strip_tags(remediationBackground)
    remediationBackground = str(remediationBackground).replace('","', ',')
    paragraph = document.add_paragraph(remediationBackground)

    document.add_heading("Additional Information:", level=3)
    # This is the fix for blank lines in CVE list
    vulnerabilityClassification = vulnerabilityClassification.split('\n')
    for i in vulnerabilityClassification:
        if len(i) > 5:
            document.add_paragraph(i, style='List Bullet')

    document.add_heading("Other Endpoints with the Same Vulnerability:", level=3)
    # add blank line to end of issue
    # paragraph = document.add_paragraph(' ')
    # paragraph = document.add_paragraph(' ')
    
    paragraph_format = paragraph.paragraph_format
    # formatting to keep our vulns together instead of line breaks
    paragraph_format.keep_together

def create_severity_graph(severity_counts):
    """
    Creates a bar graph showing the count of issues for each severity level.
    """
    # Extract severity levels and counts from severity_counts dictionary
    severities = list(severity_counts.keys())
    counts = list(severity_counts.values())

    # Plot the bar graph
    plt.figure(figsize=(6, 4))  # Adjust figure size
    bars = plt.bar(severities, counts, color=['red', 'orange', 'yellow', 'green'])
    plt.xlabel('Severity Level')
    plt.ylabel('Number of Issues')
    plt.title('Severity Distribution')
    plt.grid(axis='y', linestyle='--', alpha=0.7)  # Add grid lines

    # Add counts above each bar
    for bar, count in zip(bars, counts):
        plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height(), str(count), ha='center', va='bottom')

    buffer = io.BytesIO()
    plt.savefig(buffer, format='png')
    buffer.seek(0)
    plt.close()

    p = document.add_paragraph()
    r = p.add_run()
    r.add_picture(buffer)

def buildWordDocRepeat(host, path):
    paragraph = document.add_paragraph(host+path)
    paragraph_format = paragraph.paragraph_format
    # formatting to keep our vulns together instead of line breaks
    paragraph_format.keep_together

def severityCounts(xmlInFile):
    cwd = os.getcwd()
    xmlFileIn = os.path.join(cwd, xmlInFile)
    global issueList
    issueList = []
    # inputfile for the XML
    # THIS WILL BREAK IF YOU CHANGE HTML.PARSER!
    # try:
    if not os.path.isfile(xmlFileIn):
        status_logger.critical('Cant open XML! {}'.format(xmlInFile))
        exit(1)

    soup = BeautifulSoup(open(xmlInFile, 'r'), 'html.parser')
    status_logger.info('Using XML Input File {}'.format(xmlInFile))

    # pull all issue tags from XML
    issues = soup.findAll('issue')
    
    severity_counts = {'High': 0, 'Medium': 0, 'Low': 0, 'Information': 0}

    for i in issues:
        severity = i.find('severity')
        severity = str(severity)
        severity = strip_tags(severity)
        # print("SEVERITY",severity)
        if severity in severity_counts:
            severity_counts[severity] += 1

    return severity_counts

def pathCounts(xmlInFile):

    output_paths = []
    new_paths = []

    cwd = os.getcwd()
    xmlFileIn = os.path.join(cwd, xmlInFile)
    global issueList
    issueList = []
    # inputfile for the XML
    # THIS WILL BREAK IF YOU CHANGE HTML.PARSER!
    # try:
    if not os.path.isfile(xmlFileIn):
        status_logger.critical('Cant open XML! {}'.format(xmlInFile))
        exit(1)

    soup = BeautifulSoup(open(xmlInFile, 'r'), 'html.parser')
    status_logger.info('Using XML Input File {}'.format(xmlInFile))

    # pull all issue tags from XML
    path_xml = soup.findAll('path')
    pattern = r'<path><!\[CDATA\[(.*?)\]\]></path>'
    
    for xml_string in path_xml:
        # Use re.search to find the match
        match = re.search(pattern, str(xml_string))
        
        if match:
            # Extract the content inside the CDATA section and append to paths list
            new_paths.append(match.group(1))
    print("new_paths", new_paths)

    # for i in new_paths:
    #     print("single path", i)
    #     path = i.find('path')
        
    #     path = str(path)
    #     path = strip_tags(path)
        
    #     # print("SEVERITY",severity)
    #     output_paths.append(path)

    
    print("len of output_paths",output_paths)

    return new_paths


def process(xmlInFile):
    cwd = os.getcwd()
    xmlFileIn = os.path.join(cwd, xmlInFile)
    global issueList
    issueList = []
    # inputfile for the XML
    # THIS WILL BREAK IF YOU CHANGE HTML.PARSER!
    # try:
    if not os.path.isfile(xmlFileIn):
        status_logger.critical('Cant open XML! {}'.format(xmlInFile))
        exit(1)

    soup = BeautifulSoup(open(xmlInFile, 'r'), 'html.parser')
    status_logger.info('Using XML Input File {}'.format(xmlInFile))

    # pull all issue tags from XML
    issues = soup.findAll('issue')
    severity_order = {
    "High": 4,
    "Medium": 3,
    "Low": 2,
    "Information": 1
    }
    sorted_issues = sorted(issues, key=lambda x: severity_order.get(x.find('severity').text, 0), reverse=True)
    #sorted_issues = sorted(issues, key=lambda x: x.find('severity').text)  # 'ZZZ' ensures 'None' values are sorted to the end
    # print(sorted_issues)

    for i in sorted_issues:
        name = i.find('name').text
        host = i.find('host')
        ip = host['ip']
        host = host.text
        path = i.find('path').text
        location = i.find('location').text
        severity = i.find('severity').text
        confidence = i.find('confidence').text

        try: 
            issueBackground = i.find('issuebackground').text
            issueBackground = str(issueBackground)
        # have to replace commas before making csv. Replaced with | for now.
            issueBackground = strip_tags(issueBackground)
        # this was a fix for the CSV outfile. Need to rethink the order of this.
            issueBackground = issueBackground.replace(',', "|")
        except:
            issueBackground = 'BLANK'
            status_logger.error('issueBackground  is BLANK')

        try:
            remediationBackground = i.find('remediationbackground').text
            remediationBackground = str(remediationBackground)
            remediationBackground = strip_tags(remediationBackground)
            remediationBackground = remediationBackground.replace(',', '","')

        except:
            remediationBackground = 'BLANK'
            status_logger.error('Remediation Background is BLANK')

        try:
            vulnerabilityClassification = i.find('vulnerabilityclassifications').text
            vulnerabilityClassification = strip_tags(vulnerabilityClassification)
            status_logger.debug('Vuln Classification: {}'.format(vulnerabilityClassification))

            # Note: THis will clean HTML but also remove link, so we just get vuln name.
            # vulnerabilityClassification = strip_tags(vulnerabilityClassification)

        except:
            vulnerabilityClassification = 'BLANK'
            status_logger.error('Vuln Classification is BLANK')

        try:
            request = i.find('requestresponse').find('request').text
            request = base64.b64decode(request)
            request = str(request)
            request = response.replace(',', '","')

        except:
            request = 'BLANK'
            status_logger.error('Request is blank for {}'.format(request))

        try:
            response = i.find('requestresponse').find('response').text
            response = base64.b64decode(response)
            response = str(response)
            response = response.replace(',', '","')

        except:
            response = 'BLANK'
            status_logger.error('Response is blank for {}'.format(response))

        try:

            issueDetail = i.find('issuedetail').text
            # Not consistent with the comma encoding.
            issueDetail = str(issueDetail).replace(',', '","')


        except:
            issueDetail = 'BLANK'
            status_logger.error('Issue Detail is blank for {}'.format(issueDetail))

        # easier to build this once, so its standardized in logger and LIST
        issueLine = 'Processed Issue: [{}]'.format(str('{} ({})'.format(name, '{} Risk'.format(severity))))
        # Check here to see if issueLine already exists in the LIST. If it does, the issue is a repeat.

        if issueLine not in str(vulnList):
            # build our word document here
            buildWordDoc(name, severity, host, ip, path, location, issueBackground, issueDetail, remediationBackground,
                            vulnerabilityClassification)
            # after issue gets entered into word.
            vulnList.append(issueLine)
            issue_logger.info(issueLine)
            # paths.append(path)
        # logic if issue/vuln has already been reported on.
        if issueLine in str(vulnList):
                buildWordDocRepeat(host, path)
                # paths.append(path)
            # vulnList.append(issueLine)
            # issue_logger.info(issueLine)
            # issue_logger.warning('{} ({}) Risk: Has already been reported on! Skipping!!'.format(name, severity))
        
            # sendSkipped = (name, severity, host, ip, path, location, vulnerabilityClassification, confidence)
            # skippedVulnList.append(sendSkipped)

        """
        result = (name, host, ip, location, severity, confidence, issueBackground, remediationBackground,
                  vulnerabilityClassification, issueDetail, request, response)
                  ('{},{},{},{},{},{},{},{},{},{}').format
        """
        # document.add_page_break()
        result = (name, host, ip, location, severity, confidence, issueBackground, remediationBackground,
                  vulnerabilityClassification, issueDetail)
        issueList.append(result)
    

    status_logger.info('{} issues to report on'.format(len(issueList)))
    status_logger.info('Successfully Generate Data for Word Doc Creation')



def createSkippedVulnsOutput():
    """
    sendSkipped = (name, severity, host, ip, path, location, vulnerabilityClassification, confidence)

    """
    # add page break to get this appendix on new line
    # document.add_page_break()
    # document.add_heading('Additional Vulnerability Details', level=1)

    skippedVulnList.sort()
    for skippedVuln in skippedVulnList:
     skippedVuln= str(skippedVuln)
     confidence = skippedVuln.split(',')[7]
     confidence = confidence.split(')')[0]
     confidence = str(confidence).lower()
     if not confidence == 'tentative':
        skippedVuln = str(skippedVuln)
        skippedVuln = skippedVuln.replace("'", "")
        name = skippedVuln.split(',')[0]
        # stripping the first ( from the issue name in the report.
        name = name.split('(')[1]
        severity = skippedVuln.split(',')[1]
        host = skippedVuln.split(',')[2]
        ip = skippedVuln.split(',')[3]
        path = skippedVuln.split(',')[4]
        location = skippedVuln.split(',')[5]
        confidence = skippedVuln.split(',')[7]
        confidence = confidence.split(')')[0]
        location = str(location)
        orig_location = location
        loc_count = location.count('/')
        if loc_count < 2:
            # full_location = os.path.join(host, location)
            full_location = host + location
            location = full_location
        severity = str(severity)
        severity = severity + ' Risk '
        severity = severity.title()
        build_header = '{} ({})'.format(name, severity)
        status_logger.info('Creating Issue: {}'.format(build_header))
        document.add_heading(build_header, level=3)
        if 'http' in location:
            location = orig_location
        host_url = host + location
        host_url = host_url.replace(' ', '')

        table = document.add_table(rows=1, cols=2)
        # adjusted cell alignment here manually.
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Vulnerable Host:'
        hdr_cells[0].width = Inches(1.5)
        host = host.strip()
        hdr_cells[1].text = host
        hdr_cells[1].width = Inches(6)
        hdr_cells[1].left_margin = .1
        row_cells = table.add_row().cells
        row_cells[0].text = 'Vulnerable URL:'
        row_cells[0].width = Inches(1.5)
        host_url = host_url.strip()
        row_cells[1].text = host_url
        row_cells[1].width = Inches(6)
        row_cells[1].left_margin = .1
        #table.style = 'Light Grid Accent 1'


        table = document.add_table(rows=1, cols=2)
        # adjusted cell alignment here manually.
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Confidence:'
        hdr_cells[0].width = Inches(.00)
        confidence = confidence.strip()
        hdr_cells[1].text = confidence
        hdr_cells[1].width = Inches(.5)
        hdr_cells[1].left_margin = .1
        row_cells = table.add_row().cells
        row_cells[0].text = 'Path:'
        row_cells[0].width = Inches(.00)
        path = path.strip()
        row_cells[1].text = path
        row_cells[1].width = Inches(.5)
        row_cells[1].left_margin = .1
        #table.style = 'Light Grid Accent 1'

        table = document.add_table(rows=1, cols=2)
        # adjusted cell alignment here manually.
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'IP:'
        hdr_cells[0].width = Inches(.00)
        ip = ip.strip()
        hdr_cells[1].text = ip
        hdr_cells[1].width = Inches(.5)
        hdr_cells[1].left_margin = .1
        row_cells = table.add_row().cells
        row_cells[0].text = 'Logged in As:'
        row_cells[0].width = Inches(.00)
        path = path.strip()
        row_cells[1].text = 'Development User'
        row_cells[1].width = Inches(.5)
        row_cells[1].left_margin = .1


def writeCSV(csvFile):
    outfile = 'NOTSET'
    # csvFile = 'NOTSET'
    cwd = os.getcwd()
    csvFile = os.path.join(cwd, csvFile)
    status_logger.info('Saving to CSV file: {}'.format(csvFile))
    # need to fix this logic, still fires error instead of except:
    try:
        outfile = open(csvFile, "w", newline='')
    except:
        status_logger.critical('Cant open CSV outfile : {}'.format(outfile))

    status_logger.info('Writing to CSV'.format(csvFile))
    writer = csv.writer(outfile, delimiter=',')
    """
    writer.writerow(
        ["Name", "Host", "IP", "Path", "Severity", "Confidence", "Issue Background", "Remediation Background",
         "Vulnerability Classification", "Issue Details", "Request", "Response"])
    """
    writer.writerow(
        ["Name", "Host", "IP", "Path", "Severity", "Confidence", "Issue Background", "Remediation Background",
         "Vulnerability Classification", "Issue Details"])
    writer.writerows(issueList)


def processMultipleXmls(dir):
    if not dir:
        status_logger.critical('Supplied Dir for XML Import is blank and/or -i has not been supplied')
        exit(1)
    if not os.path.isdir(dir):
        status_logger.critical('XML Dir Specified Doesnt Exist: {}'.format(dir))
        exit(1)
    xmlList = []
    try:
        xmlList = os.listdir(dir)
    except:
        status_logger.error('The XML Process Dir must be blank : {}'.format(dir))

    for xmlFile in xmlList:
        if not xmlFile.endswith('.xml'):
            status_logger.error('Attempting to Parse non-xml file. Skipping {}'.format(xmlFile))
        if xmlFile.endswith('.xml'):
            status_logger.info('Found XML File {} in {}'.format(xmlFile, dir))
            xmlFile_path = os.path.join(dir, xmlFile)
            status_logger.debug('Processing XML file: {}'.format(xmlFile_path))
            process(xmlFile_path)

def add_item_info_to_document(items):
    """
    Adds information about the total number of items and unique items to the beginning of the document.
    """
    # Calculate total number of items and unique items
    total_items = len(items)
    unique_items = len(set(items))

    # Insert information at the beginning of the document
    p = document.add_paragraph()
    p.insert_paragraph_before("Total number of URLs accessed: ", style='Heading1').bold = True
    p.insert_paragraph_before(str(total_items))

    p = document.add_paragraph()
    p.insert_paragraph_before("Number of unique URLs accessed: ", style='Heading1').bold = True
    p.insert_paragraph_before(str(unique_items))


# def sort_issues_by_risk(document):
#     # Initialize dictionaries to store issues by risk level
#     high_risk = []
#     medium_risk = []
#     low_risk = []
#     information_risk = []

#     # Iterate through paragraphs in the document
#     for paragraph in document.paragraphs:
#         # Check if the paragraph is a heading with the specific format
#         if paragraph.style.name.startswith('Heading') and '(' in paragraph.text:
#             heading_text = paragraph.text.lower()
#             severity = heading_text.split('(')[-1].split(')')[0].strip()
#             # Sort issues based on severity
#             if severity.lower() == 'high risk':
#                 high_risk.append(paragraph)
#             elif severity.lower() == 'medium risk':
#                 medium_risk.append(paragraph)
#             elif severity.lower() == 'low risk':
#                 low_risk.append(paragraph)
#             elif severity.lower() == 'information risk':
#                 information_risk.append(paragraph)

#     # Create a new document to store sorted issues
#     sorted_document = Document()
#     # Add issues to the new document in the desired order
#     for issues in [high_risk, medium_risk, low_risk, information_risk]:
#         for issue in issues:
#             document.add_paragraph(issue.text, style=issue.style.name)
    
#     return document




def main():
    parser = optparse.OptionParser()
    parser.add_option('-i', '--xml-inputFile', help='*[REQUIRED Unless Using -d]: Specify XML Input File',
                      dest='xml_inputFile')
    parser.add_option('-o', '--word-outputFile', help='*[REQUIRED]: Specify WORD .Doc/Docx Output File',
                      dest='doc_outputFile')
    parser.add_option('-c', '--csv-outputFile', help='*[REQUIRED]: Specify CSV Output File',
                      dest='csv_outputFile')
    parser.add_option('-d', '--xml-directoryImport', help='[OPTIONAL]: Provide just a DIR to process all xml files',
                      dest='xml_processDir')
    (options, args) = parser.parse_args()
    cli_XMLFILE = options.xml_inputFile
    cli_WORDFILE = options.doc_outputFile
    cli_CSVFILE = options.csv_outputFile
    cli_XMLPROCESSDIR = options.xml_processDir

    # cli_XMLFILE =  sys.argv[1]
    xmlFileIn = cli_XMLFILE
    docOutFile = cli_WORDFILE

    if not cli_XMLFILE or cli_XMLFILE == 'None':
        if not cli_XMLPROCESSDIR:
            status_logger.critical('INPUT XML FILE NOT FOUND OR SUPPLIED!. Use -i xmlFile.xml ')
            exit(1)
    if cli_XMLFILE and cli_XMLPROCESSDIR:
        status_logger.critical('INPUT XML FILE SET AND INPUT DIR SET. Choose only 1 Bruh! ')
        exit(1)
    if not cli_WORDFILE or cli_WORDFILE == 'None' or '.doc' not in cli_WORDFILE:
        status_logger.critical('OUTPUT WORD FILE NOT FOUND OR SUPPLIED.')
        status_logger.critical('Doc/Docx Format! Use -o word.doc{word.docx}')
        exit(1)
    if not cli_CSVFILE or cli_CSVFILE == 'None' or '.csv' not in cli_CSVFILE:
        status_logger.critical('OUTPUT CSV FILE NOT FOUND OR SUPPLIED.')
        status_logger.critical('CSV Format! Use -c outFile.csv')
        exit(1)
    if not cli_XMLPROCESSDIR:
        status_logger.info('Xml DIR Import not selected')
        status_logger.debug('XML Dir Import Cli ARG : {}'.format(cli_XMLPROCESSDIR))
        severity_counts = severityCounts(xmlFileIn)
        graph = create_severity_graph(severity_counts)
        paths = pathCounts(xmlFileIn)
        add_item_info_to_document(paths)
        doc = process(xmlFileIn)
        print(doc)
        # sort_issues_by_risk(document)
        
        
    if cli_XMLPROCESSDIR:
        status_logger.info('Xml DIR Import Selected! : {}'.format(cli_XMLPROCESSDIR))
        status_logger.debug('XML Dir Import Cli ARG : {}'.format(cli_XMLPROCESSDIR))
        processMultipleXmls(cli_XMLPROCESSDIR)
    # status_logger.info('Command line XML Input file {}'.format(options.xml_inputFile))
    # logger.info('Starting The Script {}'.format(os.path.basename(__file__)))
    status_logger.info('Starting The Script {}'.format(os.path.basename(__file__)))
    status_logger.info('Using XML Input File: {}'.format(cli_XMLFILE))
    status_logger.info('Output Word Document : {}'.format(cli_WORDFILE))
    status_logger.info('Output CSV Document : {}'.format(cli_CSVFILE))
    status_logger.debug('cli_XMLFILE is Set to {}'.format(cli_XMLFILE))
    status_logger.debug('cli_WORDFILE is Set to {}'.format(cli_WORDFILE))
    status_logger.debug('cli_XMLPROCESSDIR is Set to {}'.format(cli_XMLPROCESSDIR))
    writeCSV(cli_CSVFILE)
    # generate the appendix
    # createSkippedVulnsOutput()
    # Save Word Doc
    # document.add_picture(graph, width=Inches(5)) 
    document.save(docOutFile)
    status_logger.info('Task Has Completed')


if __name__ == '__main__':
    main()
